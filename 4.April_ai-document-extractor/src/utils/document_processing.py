import docx
import PyPDF2
import os

search_index = {}

def process_uploaded_file(file_path):
    # Function to process a document based on its file type
    file_type = file_path.split('.')[-1]
    match file_type:
        case 'pdf':
            text = extract_text_from_pdf(file_path)
            return text
        case 'docx':
            text = extract_text_from_docx(file_path)
            return text
        case 'txt':
            text = extract_text_from_txt(file_path)
            return text
        case _:
            print(f"Unsupported file type: {file_type}")

def extract_text_from_pdf(pdf_path):
    # Function to extract text from a PDF file
    reader = PyPDF2.PdfReader(pdf_path)
    text = ""
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_path):
    # Function to extract text from a DOCX file
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
    return text

def extract_text_from_txt(txt_path):
    # Function to extract text from a TXT file
    with open(txt_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return text

def handle_document_upload(file):
    # Copy the file to the docs folder
    fileType = file.filename.split('.')[-1]
    # Get the path of this file
    file_path = os.path.dirname(__file__)
    print(file_path)
    path = os.path.join(file_path, 'docs', 'uploads')
    print(path)
    # Create the docs/uploads folder if it doesn't exist
    if not os.path.exists(path):
        os.makedirs(path)
    # Define the file path in the docs folder
    file_path = os.path.join(path, file.filename)
    # Check if the file already exists in the docs folder
    if os.path.exists(file_path):
        print(f"File {file.filename} already exists. Overwriting...")
    # Save the uploaded file to the docs folder
    match fileType:
        case 'pdf':
            reader = PyPDF2.PdfReader(file)
            # save the pdf file to the file path using PyPDF2
            with open(file_path, 'wb') as f:
                writer = PyPDF2.PdfWriter()
                for page_num in range(len(reader.pages)):
                    writer.add_page(reader.pages[page_num])
                writer.write(f)
        case 'docx':
            template = docx.Document()
            this_file = docx.Document(file)
            for para in this_file.paragraphs:
                template.add_paragraph(para.text)
            for table in this_file.tables:
                # Get the number of rows and columns from the table
                rows = len(table.rows)
                cols = len(table.columns)
                # Add a table to the template with the same number of rows and columns
                new_table = template.add_table(rows=rows, cols=cols)
                # Copy the content of each cell
                for i in range(rows):
                    for j in range(cols):
                        new_table.cell(i, j).text = table.cell(i, j).text
            template.save(file_path)
        case 'txt':
            fileContents = file.readlines()
            with open(file_path, 'w', encoding='utf-8') as f:
                for line in fileContents:
                    f.write(line.decode('utf-8'))
        case _:
            print(f"Unsupported file type: {fileType}")
    return file_path

def get_uploads_folder():
    """Get the uploads folder path."""
    file_path = os.path.dirname(__file__)
    return os.path.join(file_path, 'docs', 'uploads')

def get_uploaded_documents():
    """Get list of uploaded documents."""
    uploads_folder = get_uploads_folder()
    if not os.path.exists(uploads_folder):
        return []
    
    documents = []
    for filename in os.listdir(uploads_folder):
        if os.path.isfile(os.path.join(uploads_folder, filename)):
            if any(filename.lower().endswith(ext) for ext in ['.pdf', '.docx', '.txt']):
                documents.append(filename)
    return documents

def convert_file_format_from_file_path(file_path, ai_info, output_name, config_options, output_file_type):
    """
    Convert a document from file path based on user configuration options and AI insights.
    
    Args:
        file_path: Path to the file to convert.
        ai_info: Dictionary containing AI-generated insights.
        output_name: String specifying the output file name.
        config_options: Dictionary of configuration options selected by the user.
        output_file_type: String specifying the output file type.
    
    Returns:
        str: Path to the converted file.
    """
    if not os.path.exists(file_path):
        return None

    # Construct the new file path for the converted file
    base_folder = os.path.dirname(file_path)
    path = os.path.join(base_folder, '..', 'converted')
    # Create the converted folder if it doesn't exist
    if not os.path.exists(path):
        os.makedirs(path)
        
    # Determine file extension based on output type
    if output_file_type == "same":
        file_extension = os.path.splitext(file_path)[1]
    elif output_file_type == "pdf":
        file_extension = ".pdf"
    elif output_file_type == "docx":
        file_extension = ".docx"
    elif output_file_type == "txt":
        file_extension = ".txt"
    else:
        file_extension = os.path.splitext(file_path)[1]
    
    # Define the file path in the converted folder
    converted_file_path = os.path.join(path, output_name + file_extension)
    
    # Convert based on output file type
    if file_extension.lower() == '.docx':
        convert_to_docx(file_path, converted_file_path, config_options, ai_info)
    elif file_extension.lower() == '.pdf':
        convert_to_pdf(file_path, converted_file_path, config_options, ai_info)
    elif file_extension.lower() == '.txt':
        convert_to_txt(file_path, converted_file_path, config_options, ai_info)
    else:
        raise ValueError(f"Unsupported output format: {file_extension}")

    return converted_file_path

def convert_to_docx(input_path, output_path, config_dict, ai_info):
    """Convert any supported file to DOCX format."""
    input_extension = os.path.splitext(input_path)[1].lower()
    
    if input_extension == '.docx':
        # DOCX to DOCX conversion
        convert_docx_document(input_path, output_path, config_dict, ai_info)
    elif input_extension == '.pdf':
        # PDF to DOCX conversion
        text = extract_text_from_pdf(input_path)
        create_docx_from_text(text, output_path, config_dict, ai_info)
    elif input_extension == '.txt':
        # TXT to DOCX conversion
        text = extract_text_from_txt(input_path)
        create_docx_from_text(text, output_path, config_dict, ai_info)

def convert_to_pdf(input_path, output_path, config_dict, ai_info):
    """Convert any supported file to PDF format."""
    input_extension = os.path.splitext(input_path)[1].lower()
    
    if input_extension == '.pdf':
        # PDF to PDF conversion (copy with modifications)
        import shutil
        shutil.copy2(input_path, output_path)
    else:
        # Convert to DOCX first, then to PDF (placeholder - would need additional library)
        import tempfile
        import shutil
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_docx_path = temp_file.name
            convert_to_docx(input_path, temp_docx_path, config_dict, ai_info)
            # For now, just copy the DOCX file (would need python-docx2pdf or similar)
            shutil.copy2(temp_docx_path, output_path.replace('.pdf', '.docx'))
            os.unlink(temp_docx_path)

def convert_to_txt(input_path, output_path, config_dict, ai_info):
    """Convert any supported file to TXT format."""
    input_extension = os.path.splitext(input_path)[1].lower()
    
    if input_extension == '.txt':
        # TXT to TXT conversion
        import shutil
        shutil.copy2(input_path, output_path)
    elif input_extension == '.pdf':
        # PDF to TXT conversion
        text = extract_text_from_pdf(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    elif input_extension == '.docx':
        # DOCX to TXT conversion
        text = extract_text_from_docx(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)

def create_docx_from_text(text, output_path, config_dict, ai_info):
    """Create a DOCX document from plain text."""
    doc = docx.Document()
    
    # Add text to document
    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text.strip())
    
    doc.save(output_path)

def convert_docx_document(input_path, output_path, config_dict, ai_info):
    """
    Convert a DOCX document based on configuration options and AI insights.
    """
    # Load the original document
    doc = docx.Document(input_path)
    
    # Create a new document for the converted output
    new_doc = docx.Document()
    
    # Get table handling preference
    table_handling = config_dict.get('Table Handling:', 'keep')
    
    # Copy paragraphs
    for para in doc.paragraphs:
        new_doc.add_paragraph(para.text)
    
    # Handle tables based on configuration
    if table_handling == 'keep':
        for table in doc.tables:
            rows = len(table.rows)
            cols = len(table.columns)
            new_table = new_doc.add_table(rows=rows, cols=cols)
            for i in range(rows):
                for j in range(cols):
                    new_table.cell(i, j).text = table.cell(i, j).text
    elif table_handling == 'empty':
        for table in doc.tables:
            rows = len(table.rows)
            cols = len(table.columns)
            new_table = new_doc.add_table(rows=rows, cols=cols)
            # Leave cells empty
    # If 'remove', don't add tables at all
    
    new_doc.save(output_path)