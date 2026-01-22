import PyPDF2
import docx
import os
import json
import tempfile
from datetime import datetime
from pathlib import Path

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
        self.max_file_size = 100 * 1024 * 1024  # 100MB
        
    def validate_document_file(self, file):
        """Validate uploaded document file"""
        if not file:
            return False
        
        filename = file.filename.lower()
        return '.' in filename and filename.rsplit('.', 1)[1] in self.supported_formats
    
    def extract_text_from_file(self, file_path):
        """Extract text from various document formats"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                    
            return self._clean_text(text)
            
        except Exception as e:
            print(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                text = file.read()
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters that might cause issues
        text = text.replace('\x00', '')  # Remove null bytes
        text = text.replace('\ufeff', '')  # Remove BOM
        
        # Ensure text is not empty after cleaning
        if len(text.strip()) < 10:
            raise Exception("Document appears to be empty or contains no readable text")
        
        return text.strip()
    
    def analyze_document_structure(self, text):
        """Analyze document structure and content"""
        analysis = {
            'word_count': len(text.split()),
            'character_count': len(text),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
            'estimated_reading_time': round(len(text.split()) / 200, 1),  # 200 words per minute
            'language_detected': 'en',  # Simplified - could add language detection
            'document_type': self._detect_document_type(text)
        }
        
        return analysis
    
    def _detect_document_type(self, text):
        """Detect document type based on content patterns"""
        text_lower = text.lower()
        
        # Academic papers
        if any(keyword in text_lower for keyword in ['abstract', 'methodology', 'references', 'bibliography', 'hypothesis']):
            return 'academic'
        
        # Business documents
        elif any(keyword in text_lower for keyword in ['executive summary', 'quarterly', 'revenue', 'profit', 'business plan']):
            return 'business'
        
        # Legal documents
        elif any(keyword in text_lower for keyword in ['whereas', 'hereby', 'contract', 'agreement', 'terms and conditions']):
            return 'legal'
        
        # Technical documents
        elif any(keyword in text_lower for keyword in ['algorithm', 'implementation', 'system', 'technical specification', 'api']):
            return 'technical'
        
        # Research papers
        elif any(keyword in text_lower for keyword in ['study', 'research', 'findings', 'conclusion', 'analysis']):
            return 'research'
        
        else:
            return 'general'
    
    def extract_key_sections(self, text):
        """Extract key sections from document"""
        sections = {}
        lines = text.split('\n')
        current_section = "content"
        
        # Common section headers
        section_headers = [
            'abstract', 'introduction', 'methodology', 'results', 'discussion',
            'conclusion', 'references', 'executive summary', 'overview',
            'background', 'analysis', 'recommendations', 'appendix'
        ]
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            for header in section_headers:
                if header in line_lower and len(line.strip()) < 100:
                    current_section = header
                    sections[current_section] = ""
                    break
            else:
                if current_section not in sections:
                    sections[current_section] = ""
                sections[current_section] += line + "\n"
        
        return sections
    
    def export_summary(self, document_text, summary, format_type, session_id, settings=None):
        """Export summary to specified format"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            if format_type == 'txt':
                filename = f"summary_{session_id}_{timestamp}.txt"
                filepath = os.path.join(self.summaries_folder, filename)
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(f"Document Summary\n")
                    f.write(f"{'=' * 50}\n\n")
                    f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"Session ID: {session_id}\n\n")
                    
                    if settings:
                        f.write(f"Summary Settings:\n")
                        f.write(f"- Type: {settings.get('type', 'N/A')}\n")
                        f.write(f"- Length: {settings.get('length', 'N/A')}\n")
                        f.write(f"- Tone: {settings.get('tone', 'N/A')}\n\n")
                    
                    f.write(f"Summary:\n")
                    f.write(f"{'-' * 50}\n")
                    f.write(summary)
                    f.write(f"\n\n{'=' * 50}\n")
                    f.write(f"Original Document Length: {len(document_text)} characters\n")
                    f.write(f"Summary Length: {len(summary)} characters\n")
                
                return filepath
            
            elif format_type == 'json':
                filename = f"summary_{session_id}_{timestamp}.json"
                filepath = os.path.join(self.summaries_folder, filename)
                
                export_data = {
                    'session_id': session_id,
                    'timestamp': datetime.now().isoformat(),
                    'summary': summary,
                    'document_length': len(document_text),
                    'summary_length': len(summary),
                    'settings': settings or {}
                }
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(export_data, f, indent=2, ensure_ascii=False)
                
                return filepath
            
            elif format_type == 'pdf':
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                    from reportlab.lib.enums import TA_CENTER, TA_LEFT
                except ImportError:
                    # If reportlab is not installed, create a simple text file instead
                    print("Warning: reportlab not installed. Creating TXT file instead of PDF.")
                    return self.export_summary(document_text, summary, 'txt', session_id, settings)
                
                filename = f"summary_{session_id}_{timestamp}.pdf"
                filepath = os.path.join(self.summaries_folder, filename)
                
                # Create PDF
                doc = SimpleDocTemplate(filepath, pagesize=letter)
                story = []
                styles = getSampleStyleSheet()
                
                # Title style
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor='darkblue',
                    spaceAfter=30,
                    alignment=TA_CENTER
                )
                
                # Add title
                title = Paragraph("Document Summary", title_style)
                story.append(title)
                story.append(Spacer(1, 0.2*inch))
                
                # Add metadata
                metadata_style = styles['Normal']
                story.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", metadata_style))
                story.append(Paragraph(f"<b>Session ID:</b> {session_id}", metadata_style))
                story.append(Spacer(1, 0.1*inch))
                
                if settings:
                    story.append(Paragraph(f"<b>Summary Type:</b> {settings.get('type', 'N/A')}", metadata_style))
                    story.append(Paragraph(f"<b>Length:</b> {settings.get('length', 'N/A')}", metadata_style))
                    story.append(Paragraph(f"<b>Tone:</b> {settings.get('tone', 'N/A')}", metadata_style))
                
                story.append(Spacer(1, 0.3*inch))
                
                # Add summary heading
                summary_heading = Paragraph("<b>Summary:</b>", styles['Heading2'])
                story.append(summary_heading)
                story.append(Spacer(1, 0.1*inch))
                
                # Add summary content (handle line breaks and formatting)
                summary_paragraphs = summary.split('\n')
                for para in summary_paragraphs:
                    if para.strip():
                        # Handle markdown-style formatting
                        para_text = para.replace('**', '<b>').replace('**', '</b>')
                        para_text = para_text.replace('*', '<i>').replace('*', '</i>')
                        story.append(Paragraph(para_text, styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))
                
                story.append(Spacer(1, 0.2*inch))
                
                # Add statistics
                story.append(Paragraph(f"<b>Statistics:</b>", styles['Heading3']))
                story.append(Paragraph(f"Original Document: {len(document_text)} characters", metadata_style))
                story.append(Paragraph(f"Summary: {len(summary)} characters", metadata_style))
                story.append(Paragraph(f"Compression Ratio: {(1 - len(summary)/len(document_text))*100:.1f}%", metadata_style))
                
                # Build PDF
                doc.build(story)
                
                return filepath
            
            else:
                raise Exception(f"Unsupported export format: {format_type}")
        
        except Exception as e:
            print(f"Export error: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"Error exporting summary: {str(e)}")
    
    def get_file_info(self, file_path):
        """Get detailed information about uploaded file"""
        try:
            file_stats = os.stat(file_path)
            file_info = {
                'size_bytes': file_stats.st_size,
                'size_mb': round(file_stats.st_size / (1024 * 1024), 2),
                'created': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                'extension': Path(file_path).suffix.lower(),
                'filename': Path(file_path).name
            }
            return file_info
        except Exception as e:
            print(f"Error getting file info: {str(e)}")
            return {}