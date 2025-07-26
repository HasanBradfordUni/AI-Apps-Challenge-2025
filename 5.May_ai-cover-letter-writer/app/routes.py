from flask import Blueprint, render_template, request, redirect, url_for, flash, session, jsonify, send_file
import os
from .forms import UploadForm, UserForm, CVForm, JobForm, CoverLetterForm
from .models import create_connection, create_tables, add_user, add_skill, add_education, add_experience
from .models import save_cover_letter, get_cover_letter, get_user_cover_letters, find_user_by_email, execute_query
from .utils import extract_text_from_pdf, generate_cover_letter, refine_cover_letter, extract_cv_structure, model
import io
from docx import Document
import pdfkit
from datetime import datetime

app = Blueprint('app', __name__)

# Database connection
db_path = os.path.join(os.path.dirname(__file__), 'static', 'database.db')
connection = create_connection(db_path)
create_tables(connection)

@app.route('/')
def index():
    return render_template('index.html')

def get_session_cv_text():
    """Get CV text from session or return empty string"""
    return session.get('cv_text', '')

def set_session_cv_text(text):
    """Store CV text in session"""
    session['cv_text'] = text

@app.route('/upload_cv', methods=['GET', 'POST'])
def upload_cv():
    form = UploadForm()
    if form.validate_on_submit():
        # Process uploaded files
        files = request.files.getlist('files')
        if not files:
            flash('No files uploaded', 'error')
            return redirect(url_for('app.upload_cv'))
        
        # Extract text from each file
        cv_text = ""
        for file in files:
            if file:
                cv_text += extract_text_from_pdf(file)
        
        set_session_cv_text(cv_text)
        flash('CV processed successfully!', 'success')
        return redirect(url_for('app.profile'))
    
    return render_template('upload_cv.html', form=form)

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    form = UserForm()
    if form.validate_on_submit():
        # Save user profile data
        user = find_user_by_email(connection, form.email.data)
        if user:
            flash('Email already exists. Signing you in.', 'success')
            session['user_id'] = user[0]  # Assuming the first column is the user ID
            return redirect(url_for('app.cv_details'))
        else:
            user_id = add_user(connection, form.name.data, form.email.data)
            if not user_id:
                flash('Error creating user profile', 'error')
                return redirect(url_for('app.index'))
            session['user_id'] = user_id
            flash('Profile saved successfully!', 'success')
            return redirect(url_for('app.cv_details'))
    
    return render_template('profile.html', form=form)

@app.route('/cv', methods=['GET', 'POST'])
def cv_details():
    if 'user_id' not in session:
        flash('Please create a profile first', 'warning')
        return redirect(url_for('app.profile'))
    
    if request.method == 'POST':
        # Process CV file if uploaded
        if 'cv_file' in request.files and request.files['cv_file'].filename:
            cv_text = extract_text_from_pdf(request.files['cv_file'])
            set_session_cv_text(cv_text)
        
        # Process skills from dynamically added fields
        skills = []
        for key, value in request.form.items():
            if key.startswith('skills-') and key.endswith('-skill_name'):
                index = key.split('-')[1]
                skill_name = value
                proficiency = request.form.get(f'skills-{index}-proficiency', 'beginner')
                
                if skill_name.strip():  # Only add non-empty skills
                    skills.append({
                        'skill_name': skill_name,
                        'proficiency': proficiency
                    })
        
        # Process education from dynamically added fields
        education = []
        for key, value in request.form.items():
            if key.startswith('education-') and key.endswith('-institution'):
                index = key.split('-')[1]
                institution = value
                degree = request.form.get(f'education-{index}-degree', '')
                field = request.form.get(f'education-{index}-field', '')
                start_date = request.form.get(f'education-{index}-start_date', '')
                end_date = request.form.get(f'education-{index}-end_date', '')
                
                if institution.strip() and degree.strip():  # Only add entries with required fields
                    education.append({
                        'institution': institution,
                        'degree': degree,
                        'field': field,
                        'start_date': start_date,
                        'end_date': end_date
                    })
        
        # Process experience from dynamically added fields
        experience = []
        for key, value in request.form.items():
            if key.startswith('experience-') and key.endswith('-company'):
                index = key.split('-')[1]
                company = value
                position = request.form.get(f'experience-{index}-position', '')
                exp_description = request.form.get(f'experience-{index}-exp_description', '')
                start_date = request.form.get(f'experience-{index}-start_date', '')
                end_date = request.form.get(f'experience-{index}-end_date', '')
                
                if company.strip() and position.strip():  # Only add entries with required fields
                    experience.append({
                        'company': company,
                        'position': position,
                        'exp_description': exp_description,
                        'start_date': start_date,
                        'end_date': end_date
                    })
        
        # Save to database
        for skill in skills:
            add_skill(connection, session['user_id'], skill['skill_name'], skill['proficiency'])
        
        for edu in education:
            add_education(
                connection, 
                session['user_id'], 
                edu['institution'], 
                edu['degree'], 
                edu['field'], 
                edu['start_date'], 
                edu['end_date']
            )
        
        for exp in experience:
            add_experience(
                connection, 
                session['user_id'], 
                exp['company'], 
                exp['position'], 
                exp['exp_description'], 
                exp['start_date'], 
                exp['end_date']
            )
        
        flash('CV details saved successfully!', 'success')
        return redirect(url_for('app.job_details'))
    
    # Create an empty form for GET requests
    form = CVForm()
    return render_template('cv_details.html', form=form)

@app.route('/job', methods=['GET', 'POST'])
def job_details():
    if 'user_id' not in session:
        flash('Please create a profile first', 'warning')
        return redirect(url_for('app.profile'))
    
    form = JobForm()
    if form.validate_on_submit():
        # Process job description file if uploaded
        job_description = ""
        if form.job_description_file.data:
            job_description = extract_text_from_pdf(form.job_description_file.data)
        else:
            job_description = form.job_description_text.data
        
        session['job_title'] = form.job_title.data
        session['company_name'] = form.company_name.data
        session['job_description'] = job_description
        
        return redirect(url_for('app.generate_letter'))
    
    return render_template('job_details.html', form=form)

@app.route('/generate', methods=['GET', 'POST'])
def generate_letter():
    if 'user_id' not in session or 'job_description' not in session:
        flash('Please complete all previous steps first', 'warning')
        return redirect(url_for('app.index'))
    
    job_description = session.get('job_description', '')
    
    # Get customization options from form or use defaults
    tone = request.form.get('tone', 'professional')
    focus_areas = request.form.getlist('focus_areas')
    
    # PROBLEM: cv_text might not be defined here
    # FIX: Get CV text from session
    cv_text = get_session_cv_text()
    
    # Generate cover letter with customization
    cover_letter = generate_cover_letter(
        cv_text, 
        job_description, 
        tone=tone, 
        focus_areas=focus_areas
    )
    
    # If it's a regeneration request, return JSON
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'cover_letter': cover_letter})
    
    form = CoverLetterForm(cover_letter=cover_letter)
    if form.validate_on_submit():
        # Save the final cover letter
        letter_id = save_cover_letter(
            connection,
            session['user_id'],
            session.get('job_title', ''),
            session.get('company_name', ''),
            job_description,
            form.cover_letter.data
        )
        
        return redirect(url_for('app.view_letter', letter_id=letter_id))
    
    return render_template(
        'generate_letter.html', 
        form=form, 
        tone=tone, 
        focus_areas=focus_areas
    )

@app.route('/view/<int:letter_id>')
def view_letter(letter_id):
    letter_data = get_cover_letter(connection, letter_id)
    if not letter_data:
        flash('Cover letter not found', 'error')
        return redirect(url_for('app.index'))
    
    return render_template('view_letter.html', letter=letter_data)

@app.route('/history')
def letter_history():
    if 'user_id' not in session:
        flash('Please sign in first', 'warning')
        return redirect(url_for('app.profile'))
    
    letters = get_user_cover_letters(connection, session['user_id'])
    return render_template('letter_history.html', letters=letters)

@app.route('/api/refine_letter', methods=['POST', 'GET'])
def refine_letter_api():
    data = request.json
    original_letter = data.get('original_letter', '')
    feedback = data.get('feedback', '')
    
    refined_letter = refine_cover_letter(original_letter, feedback)
    return jsonify({'refined_letter': refined_letter})

@app.route('/process_cv', methods=['POST', 'GET'])
def process_cv():
    if 'cv_file' not in request.files:
        return jsonify({'success': False, 'message': 'No file uploaded'})
    
    file = request.files['cv_file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected'})
    
    # Validate file type
    if not file.filename.lower().endswith(('.pdf', '.docx', '.doc')):
        return jsonify({
            'success': False, 
            'message': 'Unsupported file format. Please upload a PDF or Word document.'
        })
    
    try:
        # Extract text using appropriate method based on file type
        if file.filename.lower().endswith('.pdf'):
            cv_text = extract_text_from_pdf(file)
        elif file.filename.lower().endswith(('.docx', '.doc')):
            cv_text = extract_text_from_word(file)
        
        # Store in session
        set_session_cv_text(cv_text)
        
        # Log extraction success with length check
        if len(cv_text) < 50:  # Simple validation
            print(f"Warning: Extracted text is very short: {cv_text}")
            return jsonify({
                'success': False,
                'message': 'The uploaded file appears to be empty or unreadable. Please try a different file.'
            })
        
        # Extract structured data using AI
        try:
            cv_data = extract_cv_structure(cv_text)
            
            if cv_data and isinstance(cv_data, dict):
                # Make sure the required keys exist
                required_keys = ['skills', 'education', 'experience']
                for key in required_keys:
                    if key not in cv_data:
                        cv_data[key] = []
                
                return jsonify({
                    'success': True,
                    'data': cv_data
                })
            else:
                # Provide a fallback structure with helpful message
                fallback_data = {
                    'skills': [{'skill_name': '', 'proficiency': 'beginner'}],
                    'education': [{'institution': '', 'degree': '', 'field': '', 'start_date': '', 'end_date': ''}],
                    'experience': [{'company': '', 'position': '', 'exp_description': '', 'start_date': '', 'end_date': ''}]
                }
                
                return jsonify({
                    'success': True,
                    'data': fallback_data,
                    'message': 'We couldn\'t automatically extract details from your CV. Please fill in the information manually.'
                })
        except Exception as e:
            print(f"Error extracting CV structure: {str(e)}")
            # Return success but with empty structure
            return jsonify({
                'success': True,
                'data': {
                    'skills': [],
                    'education': [],
                    'experience': []
                },
                'message': 'CV text extracted but structure analysis failed. Please add details manually.'
            })
    except Exception as e:
        print(f"Exception in process_cv: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Error processing CV: {str(e)}. Please try again with a different file.'
        })

@app.route('/debug_cv_extraction', methods=['GET', 'POST'])
def debug_cv_extraction():
    if request.method == 'POST':
        if 'cv_file' not in request.files:
            return "No file uploaded"
        
        file = request.files['cv_file']
        if file.filename == '':
            return "No file selected"
        
        try:
            # Extract text
            text = extract_text_from_pdf(file)
            
            # Try to get structured data
            response = model.generate_content(f"""
            Extract structured information from this CV:
            {text[:2000]}  # First 2000 chars for brevity
            
            Return ONLY valid JSON without explanations.
            """)
            
            # Return the raw response for inspection
            return f"""
            <h3>Extracted Text (first 500 chars)</h3>
            <pre>{text[:500]}</pre>
            
            <h3>Raw Gemini Response</h3>
            <pre>{response.text}</pre>
            """
        except Exception as e:
            return f"Error: {str(e)}"
            
    return """
    <form method="post" enctype="multipart/form-data">
        <input type="file" name="cv_file">
        <button type="submit">Upload and Test</button>
    </form>
    """

@app.route('/process_text_section', methods=['POST', 'GET'])
def process_text_section():
    """Process text input for CV sections using AI"""
    try:
        # Set proper content type for JSON response
        if request.method == 'GET':
            return jsonify({
                'success': False,
                'message': 'This endpoint requires a POST request with JSON data'
            }), 405
            
        data = request.json
        if not data or 'text' not in data or 'section_type' not in data:
            return jsonify({
                'success': False,
                'message': 'Invalid request data'
            }), 400
        
        text = data['text']
        section_type = data['section_type']
        
        if not text.strip():
            return jsonify({
                'success': False,
                'message': 'No text provided'
            }), 400
        
        # Construct the prompt based on section type
        if section_type == 'skills':
            prompt = generate_skills_prompt(text)
            result_key = 'skills'
        elif section_type == 'education':
            prompt = generate_education_prompt(text)
            result_key = 'education'
        elif section_type == 'experience':
            prompt = generate_experience_prompt(text)
            result_key = 'experience'
        else:
            return jsonify({
                'success': False,
                'message': 'Invalid section type'
            }), 400
        
        print(f"Processing {section_type} with text: {text[:50]}...")
        
        # Process with AI
        response = model.generate_content(prompt)
        print(f"AI response (first 100 chars): {response.text[:100]}...")
        
        result = process_ai_response(response.text, result_key)
        if result:
            print(f"Successfully parsed result with {len(result)} items")
            return jsonify({
                'success': True,
                'data': result
            })
        else:
            # If parsing fails, return a fallback structure based on section type
            print("Failed to parse AI response, using fallback")
            if section_type == 'skills':
                fallback = [{'skill_name': 'Please add skills manually', 'proficiency': 'beginner'}]
            elif section_type == 'education':
                fallback = [{'institution': 'Add institution', 'degree': 'Add degree', 
                            'field': '', 'start_date': '', 'end_date': ''}]
            elif section_type == 'experience':
                fallback = [{'company': 'Add company', 'position': 'Add position', 
                           'exp_description': '', 'start_date': '', 'end_date': ''}]
            else:
                fallback = []
                
            return jsonify({
                'success': True,
                'data': fallback,
                'message': 'Could not extract structured data, using fallback values'
            })
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error in process_text_section: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Server error: {str(e)}'
        }), 500
    
def generate_skills_prompt(text):
    """Generate a prompt for extracting skills"""
    return f"""
    Extract skills and their proficiency levels from the following text:
    
    {text}
    
    Return ONLY valid JSON with the following structure:
    [
      {{"skill_name": "Skill Name", "proficiency": "proficiency_level"}},
      ...
    ]
    
    For proficiency levels, use one of: beginner, intermediate, advanced, expert.
    Determine the proficiency based on any mentioned experience, levels, or certificates.
    If no proficiency is specified, estimate based on context.
    """

def generate_education_prompt(text):
    """Generate a prompt for extracting education"""
    return f"""
    Extract education information from the following text:
    
    {text}
    
    Return ONLY valid JSON with the following structure:
    [
      {{
        "institution": "Institution Name",
        "degree": "Degree Name",
        "field": "Field of Study",
        "start_date": "YYYY-MM-DD", 
        "end_date": "YYYY-MM-DD"
      }},
      ...
    ]
    
    Format dates as YYYY-MM-DD. If only year is available, use YYYY-01-01.
    For ongoing education, use an empty string for end_date.
    """

def generate_experience_prompt(text):
    """Generate a prompt for extracting work experience"""
    return f"""
    Extract work experience information from the following text:
    
    {text}
    
    Return ONLY valid JSON with the following structure:
    [
      {{
        "company": "Company Name",
        "position": "Position Title",
        "exp_description": "Job Description",
        "start_date": "YYYY-MM-DD",
        "end_date": "YYYY-MM-DD"
      }},
      ...
    ]
    
    Format dates as YYYY-MM-DD. If only year is available, use YYYY-01-01.
    For current positions, use an empty string for end_date.
    """

def process_ai_response(response_text, result_key):
    """Process AI response to extract JSON data with enhanced error handling"""
    try:
        import json
        import re
        
        # Log the raw response for inspection
        print(f"Raw response to parse: {response_text[:100]}...")
        
        # Clean up the response first - remove any markdown code blocks or explanations
        cleaned_response = re.sub(r'```json|```|\n\s*```|\n\s*```json|\n\s*```|\n\s*```', '', response_text)
        cleaned_response = re.sub(r'^.*?\[', '[', cleaned_response, flags=re.DOTALL)
        cleaned_response = re.sub(r'\].*?$', ']', cleaned_response, flags=re.DOTALL)
        
        # Remove any non-JSON text before or after the JSON content
        cleaned_response = cleaned_response.strip()
        
        # First attempt: direct JSON parsing
        try:
            print("Attempting direct JSON parsing")
            if cleaned_response.startswith('[') and cleaned_response.endswith(']'):
                return json.loads(cleaned_response)
            elif cleaned_response.startswith('{') and cleaned_response.endswith('}'):
                data = json.loads(cleaned_response)
                if result_key in data:
                    return data[result_key]
                else:
                    # For Gemini responses that contain JSON without the expected key
                    # Try to intelligently extract the relevant array
                    for key, value in data.items():
                        if isinstance(value, list) and len(value) > 0:
                            return value
                    return data  # Return the whole object if no lists found
        except json.JSONDecodeError as e:
            print(f"Direct JSON parsing failed: {e}")
            # Continue to regex approaches
        
        # Second attempt: find JSON array pattern using regex
        print("Attempting regex extraction of JSON array")
        array_pattern = r'\[\s*{.*?}\s*(,\s*{.*?}\s*)*\]'
        array_match = re.search(array_pattern, cleaned_response, re.DOTALL)
        if array_match:
            try:
                return json.loads(array_match.group(0))
            except json.JSONDecodeError:
                print("Array regex extraction failed")
                # Continue to next approach
        
        # Third attempt: extract individual objects and build an array
        print("Attempting to extract individual objects")
        objects = re.findall(r'{.*?}', cleaned_response, re.DOTALL)
        if objects:
            try:
                result_array = []
                for obj_str in objects:
                    obj = json.loads(obj_str)
                    result_array.append(obj)
                return result_array
            except json.JSONDecodeError:
                print("Individual object extraction failed")
                # Continue to next approach
        
        print("All parsing attempts failed")
        return None
            
    except Exception as e:
        print(f"Error processing AI response: {str(e)}")
        return None

def format_cv_text(skills_data, education_data, experience_data):
    """Format the CV text from database data"""
    
    # Format skills section
    skills_text = "SKILLS:\n"
    for skill in skills_data:
        skills_text += f"- {skill[0]} ({skill[1]})\n"
    
    # Format education section
    education_text = "\nEDUCATION:\n"
    for edu in education_data:
        institution, degree, field, start_date, end_date = edu
        education_text += f"- {degree} in {field} from {institution}"
        if start_date or end_date:
            date_range = f" ({start_date or 'N/A'} to {end_date or 'Present'})"
            education_text += date_range
        education_text += "\n"
    
    # Format experience section
    experience_text = "\nEXPERIENCE:\n"
    for exp in experience_data:
        company, position, description, start_date, end_date = exp
        experience_text += f"- {position} at {company}"
        if start_date or end_date:
            date_range = f" ({start_date or 'N/A'} to {end_date or 'Present'})"
            experience_text += date_range
        experience_text += "\n"
        if description:
            experience_text += f"  {description}\n"
    
    # Combine all sections
    formatted_cv = skills_text + education_text + experience_text
    
    return formatted_cv

@app.route('/export/<int:letter_id>/<format>')
def export_letter(letter_id, format):
    if format not in ['docx', 'pdf', 'txt']:
        flash('Unsupported export format', 'error')
        return redirect(url_for('app.view_letter', letter_id=letter_id))
    
    letter_data = get_cover_letter(connection, letter_id)
    if not letter_data:
        flash('Cover letter not found', 'error')
        return redirect(url_for('app.letter_history'))
    
    cover_letter = letter_data['content']
    filename = f"Cover_Letter_{letter_data['job_title'].replace(' ', '_')}"
    
    if format == 'docx':
        # Create Word document
        doc = Document()
        doc.add_heading(f"Cover Letter: {letter_data['job_title']}", 0)
        doc.add_paragraph(letter_data['company_name'])
        doc.add_paragraph('')
        doc.add_paragraph(cover_letter)
        
        # Save to memory
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream, 
            as_attachment=True,
            download_name=f"{filename}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    elif format == 'pdf':
        html = f"""
        <h1>Cover Letter: {letter_data['job_title']}</h1>
        <h2>{letter_data['company_name']}</h2>
        <br>
        {cover_letter.replace('\n', '<br>')}
        """
        
        pdf = pdfkit.from_string(html, False)
        
        return send_file(
            io.BytesIO(pdf),
            as_attachment=True,
            download_name=f"{filename}.pdf",
            mimetype='application/pdf'
        )
    
    else:  # txt format
        txt_content = f"Cover Letter: {letter_data['job_title']}\n"
        txt_content += f"{letter_data['company_name']}\n\n"
        txt_content += cover_letter
        
        return send_file(
            io.BytesIO(txt_content.encode()),
            as_attachment=True,
            download_name=f"{filename}.txt",
            mimetype='text/plain'
        )

@app.route('/api/letter_feedback', methods=['POST', 'GET'])
def letter_feedback():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Please sign in first'})
    
    data = request.json
    letter_id = data.get('letter_id')
    rating = data.get('rating')
    feedback = data.get('feedback', '')
    
    if not letter_id or not rating or not isinstance(rating, int) or rating < 1 or rating > 5:
        return jsonify({'success': False, 'message': 'Invalid feedback data'})
    
    try:
        # Store feedback in database
        with connection:
            cursor = connection.cursor()
            cursor.execute(
                'INSERT INTO letter_feedback (letter_id, user_id, rating, feedback, created_at) VALUES (?, ?, ?, ?, ?)',
                (letter_id, session['user_id'], rating, feedback, datetime.now())
            )
            connection.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error saving feedback: {str(e)}")
        return jsonify({'success': False, 'message': 'Error saving feedback'})

def extract_text_from_word(file):
    """Extract text from a Word document."""
    try:
        import docx
        
        # Save file to temp location
        temp_path = os.path.join(os.path.dirname(__file__), 'temp', file.filename)
        os.makedirs(os.path.dirname(temp_path), exist_ok=True)
        file.save(temp_path)
        
        # Extract text
        doc = docx.Document(temp_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Clean up temp file
        os.remove(temp_path)
        
        return text
    except Exception as e:
        print(f"Error extracting text from Word document: {str(e)}")
        raise ValueError(f"Could not process Word document: {str(e)}")